from __future__ import annotations

import base64
import logging
import mimetypes
import os
import re
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Literal
from urllib.parse import unquote, urlparse

import httpx
from fastapi.concurrency import run_in_threadpool
from openai import (
    APIConnectionError,
    APIError,
    APITimeoutError,
    AsyncOpenAI,
    AuthenticationError,
    BadRequestError,
    NotFoundError,
    OpenAIError,
    RateLimitError,
)

from app.config import get_openai_transcription_settings, load_environment
from app.services.oss_client import (
    LOCAL_UPLOADS_DIR,
    build_delivery_url_from_stored_path,
    normalize_storage_reference,
    parse_stored_file_path,
)
from app.services.persistence import extract_upload_relative_path

try:  # pragma: no cover - import path differs across moviepy releases
    from moviepy import VideoFileClip
except ImportError:  # pragma: no cover - fallback for older releases
    try:
        from moviepy.editor import VideoFileClip  # type: ignore[no-redef]
    except ImportError:  # pragma: no cover - surfaced later as a runtime parse error
        VideoFileClip = None  # type: ignore[assignment]

try:  # pragma: no cover - import guarded for environments missing optional deps
    from PyPDF2 import PdfReader
except ImportError:  # pragma: no cover - surfaced later as a user-facing error
    PdfReader = None  # type: ignore[assignment]

try:  # pragma: no cover - import guarded for environments missing optional deps
    from docx import Document
except ImportError:  # pragma: no cover - surfaced later as a user-facing error
    Document = None  # type: ignore[assignment]

try:  # pragma: no cover - import guarded for environments missing optional deps
    import pandas as pd
except ImportError:  # pragma: no cover - surfaced later as a user-facing error
    pd = None  # type: ignore[assignment]

load_environment()

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {".txt", ".md"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
SPREADSHEET_EXTENSIONS = {".csv", ".xlsx"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".aac", ".ogg", ".flac"}
MIMO_VIDEO_EXTENSIONS = {".mp4", ".mov", ".avi", ".wmv"}
MIMO_VIDEO_MIME_TYPES = {
    "video/mp4",
    "video/quicktime",
    "video/x-msvideo",
    "video/avi",
    "video/msvideo",
    "video/x-ms-wmv",
    "video/wmv",
}
MIMO_VIDEO_MAX_FILE_BYTES = 300 * 1024 * 1024
MIMO_AUDIO_EXTENSIONS = {".mp3", ".wav", ".flac", ".m4a", ".ogg"}
MIMO_AUDIO_MIME_TYPES = {
    "audio/mpeg",
    "audio/mp3",
    "audio/wav",
    "audio/x-wav",
    "audio/wave",
    "audio/flac",
    "audio/x-flac",
    "audio/mp4",
    "audio/x-m4a",
    "audio/ogg",
}
MIMO_AUDIO_MAX_FILE_BYTES = 100 * 1024 * 1024
DEFAULT_TEXT_CHAR_LIMIT = 12000
DEFAULT_OPENAI_TRANSCRIPTION_MODEL = "whisper-1"
DEFAULT_DASHSCOPE_TRANSCRIPTION_MODEL = "qwen3-asr-flash"


class MediaParserError(RuntimeError):
    """Raised when a media asset cannot be parsed into model-ready text."""


@dataclass(slots=True)
class ResolvedLocalMaterial:
    local_path: Path
    cleanup_dir: Path | None = None

    def cleanup(self) -> None:
        if self.cleanup_dir is None:
            return
        shutil.rmtree(self.cleanup_dir, ignore_errors=True)


@dataclass(frozen=True, slots=True)
class TranscriptionRuntimeConfig:
    api_key: str
    base_url: str | None
    model: str
    mode: Literal["audio_transcriptions", "dashscope_chat_completions"]


def validate_mimo_video_material(
    reference: str,
    *,
    file_size_bytes: int | None = None,
    mime_type: str | None = None,
) -> None:
    suffix = _resolve_material_suffix(reference)
    normalized_mime_type = _normalize_mime_type(mime_type)

    if suffix not in MIMO_VIDEO_EXTENSIONS and normalized_mime_type not in MIMO_VIDEO_MIME_TYPES:
        raise MediaParserError(
            "MiMo native video understanding only supports mp4, mov, avi, and wmv files.",
        )

    if file_size_bytes is not None and file_size_bytes > MIMO_VIDEO_MAX_FILE_BYTES:
        current_size_mb = file_size_bytes / (1024 * 1024)
        raise MediaParserError(
            "MiMo native video understanding only supports videos up to 300MB. "
            f"Current material is about {current_size_mb:.1f}MB.",
        )


def validate_mimo_audio_material(
    reference: str,
    *,
    file_size_bytes: int | None = None,
    mime_type: str | None = None,
) -> None:
    suffix = _resolve_material_suffix(reference)
    normalized_mime_type = _normalize_mime_type(mime_type)

    if suffix not in MIMO_AUDIO_EXTENSIONS and normalized_mime_type not in MIMO_AUDIO_MIME_TYPES:
        raise MediaParserError(
            "MiMo native audio understanding only supports mp3, wav, flac, m4a, and ogg files.",
        )

    if file_size_bytes is not None and file_size_bytes > MIMO_AUDIO_MAX_FILE_BYTES:
        current_size_mb = file_size_bytes / (1024 * 1024)
        raise MediaParserError(
            "MiMo native audio understanding only supports audio files up to 100MB. "
            f"Current material is about {current_size_mb:.1f}MB.",
        )


def _build_http_timeout(seconds: float) -> httpx.Timeout:
    connect_timeout = min(seconds, 10.0)
    return httpx.Timeout(seconds, connect=connect_timeout)


REMOTE_DOWNLOAD_TIMEOUT = _build_http_timeout(
    float(os.getenv("MEDIA_PARSER_DOWNLOAD_TIMEOUT_SECONDS", "120")),
)
TRANSCRIPTION_TIMEOUT = _build_http_timeout(
    float(
        os.getenv(
            "MEDIA_PARSER_TRANSCRIPTION_TIMEOUT_SECONDS",
            os.getenv("LLM_TIMEOUT_SECONDS", os.getenv("OPENAI_TIMEOUT_SECONDS", "120")),
        ),
    )
)

_transcription_clients: dict[tuple[str, str | None], AsyncOpenAI] = {}
_dependency_hint_emitted = False


async def parse_document(file_path_or_url: str) -> str:
    _emit_dependency_install_hint()
    try:
        resolved = await _resolve_material_to_local(file_path_or_url)
        try:
            suffix = resolved.local_path.suffix.lower()
            if suffix in TEXT_EXTENSIONS:
                extracted = await run_in_threadpool(_read_text_document, resolved.local_path)
            elif suffix in PDF_EXTENSIONS:
                extracted = await run_in_threadpool(_extract_pdf_text, resolved.local_path)
            elif suffix in DOCX_EXTENSIONS:
                extracted = await run_in_threadpool(_extract_docx_text, resolved.local_path)
            elif suffix in SPREADSHEET_EXTENSIONS:
                extracted = await run_in_threadpool(_extract_spreadsheet_text, resolved.local_path)
            else:
                raise MediaParserError(f"Unsupported document format: {suffix or 'unknown'}")

            normalized = _normalize_extracted_text(
                extracted,
                limit=int(
                    os.getenv("MEDIA_PARSER_DOCUMENT_MAX_CHARS", str(DEFAULT_TEXT_CHAR_LIMIT)),
                ),
            )
            if not normalized:
                raise MediaParserError("The document did not contain any readable text.")
            return normalized
        finally:
            resolved.cleanup()
    except MediaParserError:
        raise
    except httpx.HTTPError as exc:
        raise MediaParserError(f"Failed to download the remote document: {exc}") from exc
    except OSError as exc:
        raise MediaParserError(f"Failed to read the document: {exc}") from exc
    except Exception as exc:  # pragma: no cover - defensive guard
        raise MediaParserError(f"Document parsing failed: {exc}") from exc


async def transcribe_video(file_path_or_url: str) -> str:
    _emit_dependency_install_hint()
    runtime_config = _resolve_transcription_runtime_config()

    try:
        resolved = await _resolve_material_to_local(file_path_or_url)
        work_dir: Path | None = None
        try:
            source_path = resolved.local_path
            suffix = source_path.suffix.lower()

            if suffix in AUDIO_EXTENSIONS:
                transcription_source = source_path
            else:
                work_dir = resolved.cleanup_dir or Path(
                    tempfile.mkdtemp(prefix="omnimedia-video-audio-"),
                )
                audio_path = work_dir / f"{source_path.stem or 'transcript'}.mp3"
                await run_in_threadpool(_extract_audio_track, source_path, audio_path)
                transcription_source = audio_path

            transcript = await _request_audio_transcription(
                transcription_source,
                runtime_config=runtime_config,
            )
            normalized = _normalize_extracted_text(
                transcript,
                limit=int(
                    os.getenv("MEDIA_PARSER_TRANSCRIPT_MAX_CHARS", str(DEFAULT_TEXT_CHAR_LIMIT)),
                ),
            )
            if not normalized:
                raise MediaParserError("No usable speech transcript was produced from the video.")
            return normalized
        finally:
            if work_dir is not None and work_dir != resolved.cleanup_dir:
                shutil.rmtree(work_dir, ignore_errors=True)
            resolved.cleanup()
    except MediaParserError:
        raise
    except httpx.HTTPError as exc:
        raise MediaParserError(f"Failed to download the remote media file: {exc}") from exc
    except OSError as exc:
        raise MediaParserError(f"Video processing failed: {exc}") from exc
    except Exception as exc:  # pragma: no cover - defensive guard
        raise MediaParserError(f"Video transcription failed: {exc}") from exc


async def _resolve_material_to_local(reference: str) -> ResolvedLocalMaterial:
    normalized_reference = normalize_storage_reference(reference)
    if normalized_reference is None:
        raise MediaParserError("The material reference is empty.")

    if normalized_reference.startswith("http://") or normalized_reference.startswith("https://"):
        return await _download_remote_material(normalized_reference)

    backend_name, object_key = parse_stored_file_path(normalized_reference)
    if backend_name == "oss":
        signed_url = build_delivery_url_from_stored_path(normalized_reference)
        return await _download_remote_material(signed_url)

    relative_path = extract_upload_relative_path(normalized_reference)
    if relative_path:
        local_upload_path = LOCAL_UPLOADS_DIR / relative_path
        if not local_upload_path.exists():
            raise MediaParserError(f"Local material is missing: {local_upload_path.name}")
        return ResolvedLocalMaterial(local_path=local_upload_path)

    candidate = Path(unquote(object_key or normalized_reference))
    if candidate.exists():
        return ResolvedLocalMaterial(local_path=candidate.resolve())

    raise MediaParserError("Unable to locate the material file for parsing.")


def _resolve_material_suffix(reference: str) -> str:
    normalized_reference = normalize_storage_reference(reference) or reference
    if normalized_reference.startswith("http://") or normalized_reference.startswith("https://"):
        candidate = urlparse(normalized_reference).path or normalized_reference
    else:
        _, object_key = parse_stored_file_path(normalized_reference)
        candidate = object_key or normalized_reference
    return Path(unquote(candidate)).suffix.lower()


def _normalize_mime_type(mime_type: str | None) -> str:
    return (mime_type or "").split(";", 1)[0].strip().lower()


async def _download_remote_material(url: str) -> ResolvedLocalMaterial:
    logger.info("Downloading remote material for parsing: %s", url)
    async with httpx.AsyncClient(
        timeout=REMOTE_DOWNLOAD_TIMEOUT,
        follow_redirects=True,
    ) as client:
        response = await client.get(url)
        response.raise_for_status()

    if not response.content:
        raise MediaParserError("Remote material download returned empty content.")

    download_dir = Path(tempfile.mkdtemp(prefix="omnimedia-material-"))
    suffix = _resolve_suffix_from_download(
        url=url,
        content_type=response.headers.get("content-type"),
    )
    file_name = _guess_filename(url) or "material"
    destination = download_dir / f"{Path(file_name).stem or 'material'}{suffix}"
    await run_in_threadpool(destination.write_bytes, response.content)
    return ResolvedLocalMaterial(local_path=destination, cleanup_dir=download_dir)


def _resolve_suffix_from_download(*, url: str, content_type: str | None) -> str:
    url_suffix = Path(urlparse(url).path).suffix.lower()
    if url_suffix:
        return url_suffix

    normalized_type = (content_type or "").split(";", 1)[0].strip().lower()
    guessed_suffix = mimetypes.guess_extension(normalized_type)
    if guessed_suffix:
        return guessed_suffix
    return ".bin"


def _guess_filename(reference: str) -> str:
    parsed = urlparse(reference)
    path_name = Path(unquote(parsed.path or reference)).name
    if path_name:
        return path_name
    return "material"


def _read_text_document(path: Path) -> str:
    raw_bytes = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def _extract_pdf_text(path: Path) -> str:
    if PdfReader is None:
        raise MediaParserError(
            "PyPDF2 is not installed. Run `pip install -r requirements.txt` first.",
        )

    reader = PdfReader(str(path))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:  # pragma: no cover - defensive guard
            raise MediaParserError("Encrypted PDF files are not supported right now.") from exc

    pages: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)

    return "\n\n".join(pages)


def _extract_docx_text(path: Path) -> str:
    if Document is None:
        raise MediaParserError(
            "python-docx is not installed. Run `pip install -r requirements.txt` first.",
        )

    document = Document(str(path))
    sections: list[str] = []

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    if paragraphs:
        sections.append("\n".join(paragraphs))

    table_blocks: list[str] = []
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [
                re.sub(r"\s*\n\s*", " / ", cell.text).strip()
                for cell in row.cells
            ]
            normalized_cells = [cell for cell in cells if cell]
            if normalized_cells:
                rows.append(" | ".join(normalized_cells))
        if rows:
            table_blocks.append("\n".join(rows))

    if table_blocks:
        sections.append("\n\n".join(table_blocks))

    return "\n\n".join(section for section in sections if section.strip())


def _extract_spreadsheet_text(path: Path) -> str:
    if pd is None:
        raise MediaParserError(
            "pandas and openpyxl are not installed. Run `pip install -r requirements.txt` first.",
        )

    try:
        if path.suffix.lower() == ".csv":
            workbook: dict[str, object] = {
                "Sheet1": pd.read_csv(  # type: ignore[union-attr]
                    str(path),
                    dtype=str,
                    keep_default_na=False,
                )
            }
        else:
            raw_workbook = pd.read_excel(  # type: ignore[union-attr]
                str(path),
                sheet_name=None,
                dtype=str,
                keep_default_na=False,
            )
            if isinstance(raw_workbook, dict):
                workbook = {
                    str(sheet_name).strip() or f"Sheet {index + 1}": frame
                    for index, (sheet_name, frame) in enumerate(raw_workbook.items())
                }
            else:
                workbook = {"Sheet1": raw_workbook}
    except Exception as exc:
        logger.error("Spreadsheet parsing failed for %s: %s", path.name, exc)
        raise MediaParserError(
            "Failed to parse the spreadsheet. Ensure the file is a valid CSV or Excel workbook.",
        ) from exc

    row_blocks: list[str] = []
    include_sheet_name = path.suffix.lower() == ".xlsx" or len(workbook) > 1
    for sheet_name, frame in workbook.items():
        if frame is None:
            continue
        normalized_frame = frame.fillna("")
        columns = [
            _normalize_spreadsheet_header(raw_header, index)
            for index, raw_header in enumerate(getattr(normalized_frame, "columns", []))
        ]
        for row_index, row in enumerate(
            normalized_frame.itertuples(index=False, name=None),
            start=1,
        ):
            pairs: list[str] = []
            for column_index, raw_value in enumerate(row):
                value = _normalize_spreadsheet_value(raw_value)
                if not value:
                    continue
                header = (
                    columns[column_index]
                    if column_index < len(columns)
                    else f"Column {column_index + 1}"
                )
                pairs.append(f"{header}: {value}")
            if not pairs:
                continue

            prefix_parts: list[str] = []
            if include_sheet_name:
                prefix_parts.append(f"Sheet: {sheet_name}")
            prefix_parts.append(f"Row: {row_index}")
            row_blocks.append(f"{' | '.join(prefix_parts)} | {', '.join(pairs)}")

    if not row_blocks:
        raise MediaParserError("The spreadsheet did not contain any readable rows.")

    return "\n".join(row_blocks)


def _normalize_spreadsheet_header(raw_header: object, index: int) -> str:
    normalized = re.sub(r"\s+", " ", str(raw_header or "")).strip()
    return normalized or f"Column {index + 1}"


def _normalize_spreadsheet_value(raw_value: object) -> str:
    normalized = re.sub(r"\s+", " ", str(raw_value or "")).strip()
    if normalized.lower() in {"", "nan", "nat", "none"}:
        return ""
    return normalized


def _extract_audio_track(source_path: Path, audio_path: Path) -> None:
    if VideoFileClip is None:
        raise MediaParserError(
            "moviepy is not installed. Run `pip install -r requirements.txt` first.",
        )

    clip = VideoFileClip(str(source_path))
    audio_clip = clip.audio
    try:
        if audio_clip is None:
            raise MediaParserError("The uploaded video does not contain an audio track.")
        audio_clip.write_audiofile(
            str(audio_path),
            codec="libmp3lame",
            bitrate="64k",
            ffmpeg_params=["-ac", "1", "-ar", "16000"],
            logger=None,
        )
    except OSError as exc:
        raise MediaParserError(
            "Audio extraction failed. Ensure ffmpeg is installed and available to moviepy.",
        ) from exc
    finally:
        if audio_clip is not None:
            audio_clip.close()
        clip.close()


async def _request_audio_transcription(
    audio_path: Path,
    *,
    runtime_config: TranscriptionRuntimeConfig,
) -> str:
    if runtime_config.mode == "dashscope_chat_completions":
        return await _request_dashscope_compatible_transcription(
            audio_path,
            runtime_config=runtime_config,
        )
    return await _request_openai_audio_transcription(
        audio_path,
        runtime_config=runtime_config,
    )


async def _request_openai_audio_transcription(
    audio_path: Path,
    *,
    runtime_config: TranscriptionRuntimeConfig,
) -> str:
    try:
        with audio_path.open("rb") as audio_file:
            response = await _get_transcription_client(runtime_config).audio.transcriptions.create(
                model=runtime_config.model,
                file=audio_file,
            )
    except AuthenticationError as exc:
        raise MediaParserError("Audio transcription authentication failed.") from exc
    except RateLimitError as exc:
        raise MediaParserError("Audio transcription hit a rate limit.") from exc
    except APITimeoutError as exc:
        raise MediaParserError("Audio transcription timed out.") from exc
    except APIConnectionError as exc:
        raise MediaParserError("Could not connect to the audio transcription service.") from exc
    except NotFoundError as exc:
        raise MediaParserError(
            "The configured transcription gateway does not expose /audio/transcriptions. "
            "Set OPENAI_TRANSCRIPTION_BASE_URL or LLM_BASE_URL to a provider that supports it.",
        ) from exc
    except BadRequestError as exc:
        raise MediaParserError(f"Audio transcription request was rejected: {exc}") from exc
    except (APIError, OpenAIError) as exc:
        raise MediaParserError(f"Audio transcription request failed: {exc}") from exc

    transcript = getattr(response, "text", "")
    if isinstance(transcript, str):
        return transcript

    if isinstance(response, dict):
        return str(response.get("text", ""))
    return str(transcript)


async def _request_dashscope_compatible_transcription(
    audio_path: Path,
    *,
    runtime_config: TranscriptionRuntimeConfig,
) -> str:
    data_uri = await run_in_threadpool(_build_audio_data_uri, audio_path)

    try:
        response = await _get_transcription_client(runtime_config).chat.completions.create(
            model=runtime_config.model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_audio",
                            "input_audio": {
                                "data": data_uri,
                            },
                        }
                    ],
                }
            ],
            stream=False,
            extra_body={
                "asr_options": {
                    "enable_itn": False,
                }
            },
        )
    except AuthenticationError as exc:
        raise MediaParserError("DashScope-compatible transcription authentication failed.") from exc
    except RateLimitError as exc:
        raise MediaParserError("DashScope-compatible transcription hit a rate limit.") from exc
    except APITimeoutError as exc:
        raise MediaParserError("DashScope-compatible transcription timed out.") from exc
    except APIConnectionError as exc:
        raise MediaParserError("Could not connect to the DashScope-compatible transcription service.") from exc
    except BadRequestError as exc:
        raise MediaParserError(
            "DashScope-compatible transcription rejected the request. "
            "Check audio duration, size, and model availability.",
        ) from exc
    except (APIError, OpenAIError) as exc:
        raise MediaParserError(f"DashScope-compatible transcription failed: {exc}") from exc

    message_content = response.choices[0].message.content if response.choices else ""
    normalized = _coerce_chat_message_text(message_content)
    if normalized.strip():
        return normalized
    raise MediaParserError("DashScope-compatible transcription returned an empty response.")


def _get_transcription_client(
    runtime_config: TranscriptionRuntimeConfig,
) -> AsyncOpenAI:
    cache_key = (runtime_config.api_key, runtime_config.base_url)
    if cache_key in _transcription_clients:
        return _transcription_clients[cache_key]

    client_kwargs: dict[str, object] = {
        "api_key": runtime_config.api_key,
        "timeout": TRANSCRIPTION_TIMEOUT,
    }
    if runtime_config.base_url:
        client_kwargs["base_url"] = runtime_config.base_url
    client = AsyncOpenAI(**client_kwargs)
    _transcription_clients[cache_key] = client
    return client


def _resolve_transcription_runtime_config() -> TranscriptionRuntimeConfig:
    settings = get_openai_transcription_settings()
    if settings is None or not settings.api_key:
        raise MediaParserError("No transcription API key is configured.")

    return _build_transcription_runtime_config(
        api_key=settings.api_key,
        base_url=settings.base_url or None,
        model_override=settings.model or None,
    )


def _build_transcription_runtime_config(
    *,
    api_key: str,
    base_url: str | None,
    model_override: str | None = None,
) -> TranscriptionRuntimeConfig:
    if _is_dashscope_compatible_base_url(base_url):
        return TranscriptionRuntimeConfig(
            api_key=api_key,
            base_url=base_url,
            model=(model_override or "").strip() or DEFAULT_DASHSCOPE_TRANSCRIPTION_MODEL,
            mode="dashscope_chat_completions",
        )

    return TranscriptionRuntimeConfig(
        api_key=api_key,
        base_url=base_url,
        model=(model_override or "").strip() or DEFAULT_OPENAI_TRANSCRIPTION_MODEL,
        mode="audio_transcriptions",
    )


def _is_dashscope_compatible_base_url(base_url: str | None) -> bool:
    normalized = (base_url or "").strip().lower()
    return "dashscope" in normalized and "/compatible-mode/" in normalized


def _build_audio_data_uri(audio_path: Path) -> str:
    mime_type = _resolve_audio_mime_type(audio_path)
    encoded = base64.b64encode(audio_path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _resolve_audio_mime_type(audio_path: Path) -> str:
    guessed = mimetypes.guess_type(audio_path.name)[0]
    if guessed:
        return guessed

    suffix = audio_path.suffix.lower()
    if suffix == ".mp3":
        return "audio/mpeg"
    if suffix == ".wav":
        return "audio/wav"
    return "application/octet-stream"


def _coerce_chat_message_text(message_content: object) -> str:
    if isinstance(message_content, str):
        return message_content
    if isinstance(message_content, list):
        parts: list[str] = []
        for item in message_content:
            if isinstance(item, str):
                parts.append(item)
                continue
            if isinstance(item, dict):
                if isinstance(item.get("text"), str):
                    parts.append(str(item["text"]))
                    continue
                if item.get("type") == "text" and isinstance(item.get("text"), str):
                    parts.append(str(item["text"]))
        return "\n".join(part.strip() for part in parts if part.strip())
    return str(message_content or "")


def _normalize_extracted_text(raw_text: str, *, limit: int) -> str:
    normalized = raw_text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    normalized = "\n".join(line.rstrip() for line in normalized.split("\n"))
    normalized = normalized.strip()
    if not normalized:
        return ""

    if len(normalized) <= limit:
        return normalized

    truncated = normalized[:limit].rstrip()
    return f"{truncated}\n\n[Content truncated after {limit} characters.]"


def _emit_dependency_install_hint() -> None:
    global _dependency_hint_emitted

    if _dependency_hint_emitted:
        return

    logger.warning(
        "Media parser reminder: run `pip install moviepy PyPDF2 python-docx pandas openpyxl openai` and ensure ffmpeg is available before debugging long-form document, spreadsheet, or video parsing.",
    )
    _dependency_hint_emitted = True
